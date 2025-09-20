from docx import Document  # type: ignore
import csv
import os
import io
from typing import Dict, List, Set, Optional, Callable, Tuple, Any
from openpyxl import load_workbook  # type: ignore
from datetime import date, datetime

def replace_text_in_paragraph(paragraph: Any, replacements: dict) -> None:
    """Replace multiple placeholders in a paragraph at once."""
    full_text = ''.join(run.text for run in paragraph.runs)
    modified = False
    
    for placeholder, replacement in replacements.items():
        if placeholder in full_text:
            full_text = full_text.replace(placeholder, replacement)
            modified = True
    
    if modified:
        # Clear existing runs
        for run in paragraph.runs:
            run.text = ''
        # Set the new text as a single run
        if paragraph.runs:
            paragraph.runs[0].text = full_text
        else:
            paragraph.add_run(full_text)

def replace_placeholders(doc: Document, replacements: dict) -> None:
    """Replace all placeholders in the document at once."""
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, replacements)

 


def generate_placeholder_variants(placeholder: str) -> Set[str]:
    """Generate reasonable variants for a placeholder to improve matching robustness."""
    variants: Set[str] = set()
    raw = placeholder
    variants.add(raw)

    # Ensure closing bracket if missing
    if raw.startswith('[') and not raw.endswith(']'):
        variants.add(raw + ']')

    # Normalize en dash and hyphen both ways
    if '–' in raw:
        variants.add(raw.replace('–', '-'))
    if '-' in raw:
        variants.add(raw.replace('-', '–'))

    return variants


DECLARED_PLACEHOLDERS: List[str] = [
    "[Programme Title]",
    "[Learner Registration Number]",
    "[Learner Name]",
    "[Assignment Title]",
    "[Assessor Name]",
    "[Unit/Component Number and Title]",
    "[Targeted Learning Aims/Assessment Criteria (Initial)]",
    "[First Submission - Deadline]",
    "[First Submission - Date Submitted]",
    "[Extension Approved (Y/N)]",
    "[Initial - General Comments]",
    "[Initial - Learner Signature (Name or File Path)]",
    "[Initial - Learner Declaration Date]",
    "[Initial - Assessor Signature (Name or File Path)]",
    "[Initial - Assessor Declaration Date]",
    "[Initial - Date of Feedback to Learner]",
    "[Resubmission - Authorised by Lead Internal Verifier (Name)]",
    "[Resubmission - Authorisation Date]",
    "[Resubmission - Deadline]",
    "[Resubmission - Date Submitted]",
    "[Resubmission - General Comments]",
    "[Resubmission - Learner Signature (Name or File Path)]",
    "[Resubmission - Learner Declaration Date]",
    "[Resubmission - Assessor Signature (Name or File Path)]",
    "[Resubmission - Assessor Declaration Date]",
    "[Resubmission - Date of Feedback to Learner]",
    "[Retake - Deadline]",
    "[Retake - Date Submitted]",
    "[Retake - General Comments]",
    "[Retake - Learner Signature (Name or File Path)]",
    "[Retake - Learner Declaration Date]",
    "[Retake - Assessor Signature (Name or File Path)]",
    "[Retake - Assessor Declaration Date]",
    "[Retake - Date of Feedback to Learner]",
]


def process_criteria(targeted: str, achieved: str, max_criteria: int = 3) -> tuple[list[str], list[str]]:
    """Process targeted and achieved criteria into lists and Y/N markers.
    
    Args:
        targeted: Comma-separated string of targeted criteria
        achieved: Comma-separated string of achieved criteria
        max_criteria: Maximum number of criteria to process (default 3)
    """
    # Split and clean targeted criteria
    targeted_list = [c.strip() for c in targeted.split(',') if c.strip()]
    # Split and clean achieved criteria
    achieved_list = [c.strip() for c in achieved.split(',') if c.strip()]
    
    # Generate Y/N list based on whether each targeted criteria was achieved
    achieved_yn = ['Y' if t in achieved_list else 'N' for t in targeted_list]
    
    # Pad both lists to specified length with empty strings
    targeted_list.extend([''] * (max_criteria - len(targeted_list)))
    achieved_yn.extend([''] * (max_criteria - len(achieved_yn)))
    
    return targeted_list[:max_criteria], achieved_yn[:max_criteria]

def replace_all_placeholders(doc: Document, row: Dict[str, str]) -> None:
    """Replace placeholders in doc using both declared list and dynamic [Header] placeholders."""
    # Build mapping: placeholder variant -> replacement value
    replacement_map: Dict[str, str] = {}

    # Handle Initial criteria placeholders (up to 3)
    initial_targeted = row.get('Initial - Targeted Criteria', '').strip()
    initial_achieved = row.get('Initial - Criteria Achieved', '').strip()
    initial_targeted_list, initial_achieved_yn = process_criteria(initial_targeted, initial_achieved, max_criteria=3)
    
    # Add Initial criteria mappings
    for i, (target, achieved) in enumerate(zip(initial_targeted_list, initial_achieved_yn), 1):
        replacement_map[f'[ITC{i}]'] = target
        replacement_map[f'[ICA{i}]'] = achieved

    # Handle Resubmission criteria placeholders (up to 5)
    resub_targeted = row.get('Resubmission - Targeted Criteria', '').strip()
    resub_achieved = row.get('Resubmission - Criteria Achieved', '').strip()
    resub_targeted_list, resub_achieved_yn = process_criteria(resub_targeted, resub_achieved, max_criteria=5)
    
    # Add Resubmission criteria mappings
    for i, (target, achieved) in enumerate(zip(resub_targeted_list, resub_achieved_yn), 1):
        replacement_map[f'[RTC{i}]'] = target
        replacement_map[f'[RCA{i}]'] = achieved

    # Declared placeholders from specification
    for placeholder in DECLARED_PLACEHOLDERS:
        column_name = placeholder.strip().lstrip('[').rstrip(']')
        value = (row.get(column_name) or '').strip()
        for variant in generate_placeholder_variants(placeholder):
            replacement_map[variant] = value

    # Perform all replacements at once
    replace_placeholders(doc, replacement_map)


_template_cache = {}

def get_template_doc(template_path: str) -> Document:
    """Get a cached template document or create a new one."""
    if template_path not in _template_cache:
        _template_cache[template_path] = Document(template_path)
    return Document(template_path)  # Return a fresh copy from the template

from concurrent.futures import ProcessPoolExecutor, as_completed
import multiprocessing

def _process_single_row(args) -> Tuple[int, Tuple[Optional[str], Any]]:
    """Process a single row and return the generated document.
    
    Returns:
        Tuple containing (index, (filename, doc_bytes)) on success
        or (index, (None, error_message)) on failure
    """
    index, row, template_path = args
    try:
        # Create a fresh document from the template
        doc = Document(template_path)  # Don't use cache in worker processes
        replace_all_placeholders(doc, row)

        # Create simple filename from name and registration number
        name = (row.get('Learner Name') or '').strip()
        reg = (row.get('Learner Registration Number') or '').strip()
        filename = f"{name} {reg}.docx".strip()

        # Save document to bytes buffer
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_bytes = doc_buffer.getvalue()
        doc_buffer.close()

        return index, (filename, doc_bytes)
    except (IOError, ValueError, KeyError) as e:
        return index, (None, f"Error processing row: {str(e)}")
    except Exception as e:
        return index, (None, f"Unexpected error: {str(e)}")

def generate_documents_from_csv(
    csv_path: str,
    template_path: str,
    progress: Optional[Callable[[str, Dict[str, object]], None]] = None,
    max_workers: Optional[int] = None
) -> List[Tuple[str, bytes]]:
    """Generate documents in memory and return list of (filename, document_bytes) tuples."""
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found: {template_path}")
    if not os.path.exists(csv_path):
        raise FileNotFoundError(f"CSV not found: {csv_path}")

    # Pre-count total data rows for progress reporting
    try:
        with open(csv_path, mode='r', encoding='utf-8-sig', newline='') as f_count:
            total_rows = max(0, sum(1 for _ in csv.reader(f_count)) - 1)
    except Exception:
        total_rows = 0

    if progress:
        progress('start', {'total_rows': total_rows})

    # Read all rows into memory
    with open(csv_path, mode='r', encoding='utf-8-sig', newline='') as f:
        rows = list(csv.DictReader(f))

    generated_docs = []

    # Determine number of workers
    if max_workers is None:
        # Leave one core free for the main process
        max_workers = max(1, multiprocessing.cpu_count() - 1)

    # Create work items
    work_items = [(i, row, template_path) for i, row in enumerate(rows)]

    with ProcessPoolExecutor(max_workers=max_workers) as executor:
        # Submit all work items
        future_to_index = {
            executor.submit(_process_single_row, item): item[0]
            for item in work_items
        }

        # Process results as they complete
        for future in as_completed(future_to_index):
            index = future_to_index[future]
            try:
                index, result = future.result()
                filename, doc_bytes = result

                if filename is None:  # Error occurred
                    if progress:
                        progress('row_error', {'index': index, 'error': doc_bytes})  # doc_bytes contains error message
                    continue

                generated_docs.append((filename, doc_bytes))
                
                if progress:
                    progress('row_done', {'index': index, 'filename': filename})

            except Exception as e:
                if progress:
                    progress('row_error', {'index': index, 'error': str(e)})

    if progress:
        progress('complete', {'generated': len(generated_docs), 'total_rows': total_rows})

    return generated_docs

def convert_xlsx_to_csv(xlsx_path: str, csv_path: str) -> None:
    """Convert the first worksheet of an .xlsx file to a UTF-8 CSV file."""
    if not os.path.exists(xlsx_path):
        raise FileNotFoundError(f"XLSX not found: {xlsx_path}")

    # Use read_only mode and data_only for better performance
    wb = load_workbook(filename=xlsx_path, read_only=True, data_only=True)
    ws = wb.worksheets[0]

    # Pre-compile the format function for better performance
    def format_cell_for_csv(value):
        if value is None:
            return ""
        if isinstance(value, (datetime, date)):
            return value.isoformat().split('T')[0]
        return str(value)

    # Use a buffer to write rows in batches
    buffer_size = 1000
    row_buffer = []

    with open(csv_path, mode='w', encoding='utf-8', newline='') as f:
        writer = csv.writer(f)
        
        for row in ws.iter_rows(values_only=True):
            row_buffer.append([format_cell_for_csv(v) for v in row])
            
            if len(row_buffer) >= buffer_size:
                writer.writerows(row_buffer)
                row_buffer = []
        
        # Write any remaining rows
        if row_buffer:
            writer.writerows(row_buffer)

    wb.close()


if __name__ == "__main__":
    source_dir = os.path.dirname(os.path.abspath(__file__))
    TEMPLATE_PATH = os.path.join(source_dir, "template.docx")
    XLSX_PATH = os.path.join(source_dir, "dummy_data.xlsx")

    # Convert XLSX to a temp CSV file
    import tempfile
    with tempfile.TemporaryDirectory() as tmpdir:
        temp_csv_path = os.path.join(tmpdir, "dummy.csv")
        convert_xlsx_to_csv(XLSX_PATH, temp_csv_path)
        docs = generate_documents_from_csv(temp_csv_path, TEMPLATE_PATH)
        print(f"Generated {len(docs)} documents")

